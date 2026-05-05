#!/usr/bin/env python3
"""
Универсальный скрипт для извлечения текста, конвертации документов и OCR.

Использование:
    python extract_text.py <путь_к_файлу> [опции]

Опции:
    --tables      Извлекать таблицы (PDF/XLSX)
    --metadata    Показать метаданные
    --output F    Сохранить результат в файл
    --json        Вывод в формате JSON
    --convert     Конвертировать файл
    --to F        Формат для конвертации (pdf, docx, md, html)
    --ocr         Распознать текст (OCR)
    --lang L      Язык OCR (rus, eng, rus+eng)
    --help        Показать справку
"""

import argparse
import json
import os
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

# Цвета для консоли
class Colors:
    HEADER = '\033[95m'
    BLUE = '\033[94m'
    CYAN = '\033[96m'
    GREEN = '\033[92m'
    YELLOW = '\033[93m'
    RED = '\033[91m'
    ENDC = '\033[0m'
    BOLD = '\033[1m'


def print_info(msg):
    print(f"{Colors.CYAN}[INFO]{Colors.ENDC} {msg}")


def print_error(msg):
    print(f"{Colors.RED}[ERROR]{Colors.ENDC} {msg}")


def print_success(msg):
    print(f"{Colors.GREEN}[OK]{Colors.ENDC} {msg}")


def print_warning(msg):
    print(f"{Colors.YELLOW}[WARN]{Colors.ENDC} {msg}")


def detect_format(file_path: str) -> str:
    """Определить формат файла по расширению."""
    ext = Path(file_path).suffix.lower()
    format_map = {
        '.pdf': 'pdf',
        '.docx': 'docx',
        '.doc': 'docx',
        '.xlsx': 'xlsx',
        '.xls': 'xlsx',
        '.txt': 'txt',
    }
    return format_map.get(ext, 'unknown')


def extract_pdf(file_path: str, extract_tables: bool = False, show_metadata: bool = False) -> dict:
    """Извлечь текст из PDF."""
    try:
        from PyPDF2 import PdfReader
        import pdfplumber
    except ImportError:
        return {"error": "Установите библиотеки: pip install PyPDF2 pdfplumber"}

    result = {"format": "pdf", "pages": []}

    # Метаданные
    if show_metadata:
        reader = PdfReader(file_path)
        result["metadata"] = {
            "pages": len(reader.pages),
            "title": getattr(reader.metadata, 'title', None),
            "author": getattr(reader.metadata, 'author', None),
            "subject": getattr(reader.metadata, 'subject', None),
            "creator": getattr(reader.metadata, 'creator', None),
        }

    # Текст и таблицы
    with pdfplumber.open(file_path) as pdf:
        result["page_count"] = len(pdf.pages)

        for i, page in enumerate(pdf.pages):
            page_data = {"page": i + 1}

            # Текст
            text = page.extract_text()
            if text:
                page_data["text"] = text.strip()

            # Таблицы
            if extract_tables:
                tables = page.extract_tables()
                if tables:
                    page_data["tables"] = []
                    for j, table in enumerate(tables):
                        if table:
                            page_data["tables"].append({
                                "table_num": j + 1,
                                "rows": len(table),
                                "data": table
                            })

            result["pages"].append(page_data)

    return result


def extract_docx(file_path: str, show_metadata: bool = False) -> dict:
    """Извлечь текст из DOCX."""
    try:
        from docx import Document
    except ImportError:
        return {"error": "Установите библиотеку: pip install python-docx"}

    doc = Document(file_path)
    result = {
        "format": "docx",
        "paragraphs": [],
        "tables": [],
    }

    # Метаданные
    if show_metadata:
        props = doc.core_properties
        result["metadata"] = {
            "author": props.author,
            "title": props.title,
            "subject": props.subject,
            "keywords": props.keywords,
        }

    # Параграфы
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            result["paragraphs"].append({
                "num": i + 1,
                "style": para.style.name if para.style else None,
                "text": para.text.strip()
            })

    # Таблицы
    for i, table in enumerate(doc.tables):
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)

        if table_data:
            result["tables"].append({
                "num": i + 1,
                "rows": len(table_data),
                "cols": len(table_data[0]) if table_data else 0,
                "data": table_data
            })

    return result


def extract_xlsx(file_path: str, extract_tables: bool = False, show_metadata: bool = False) -> dict:
    """Извлечь данные из XLSX."""
    try:
        from openpyxl import load_workbook
        import pandas as pd
    except ImportError:
        return {"error": "Установите библиотеки: pip install openpyxl pandas"}

    result = {
        "format": "xlsx",
        "sheets": [],
    }

    wb = load_workbook(file_path, data_only=True)
    result["sheet_names"] = wb.sheetnames
    result["sheet_count"] = len(wb.sheetnames)

    # Метаданные
    if show_metadata:
        props = wb.properties
        result["metadata"] = {
            "title": props.title,
            "author": props.creator,
            "subject": props.subject,
            "created": str(props.created) if props.created else None,
        }

    # Данные каждого листа
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        sheet_data = {
            "name": sheet_name,
            "max_row": sheet.max_row,
            "max_col": sheet.max_column,
            "data": []
        }

        # Первые 100 строк для предпросмотра
        rows_shown = 0
        max_preview = 100

        for row in sheet.iter_rows(max_row=max_preview, values_only=True):
            if any(cell is not None for cell in row):
                sheet_data["data"].append(list(row))
                rows_shown += 1

        if rows_shown < sheet.max_row:
            sheet_data["truncated"] = True
            sheet_data["total_rows"] = sheet.max_row
            sheet_data["preview_rows"] = rows_shown

        result["sheets"].append(sheet_data)

    # Таблицы (если запрошено) - читаем через pandas
    if extract_tables:
        try:
            xls = pd.ExcelFile(file_path)
            result["tables"] = []
            for sheet_name in xls.sheet_names:
                df = pd.read_excel(xls, sheet_name=sheet_name)
                if not df.empty:
                    result["tables"].append({
                        "sheet": sheet_name,
                        "rows": len(df),
                        "columns": list(df.columns),
                        "data": df.head(50).to_dict('records')
                    })
        except Exception:
            pass

    wb.close()
    return result


def extract_txt(file_path: str, show_metadata: bool = False) -> dict:
    """Извлечь текст из TXT."""
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()

    result = {
        "format": "txt",
        "text": content,
        "lines": len(content.splitlines()),
        "chars": len(content)
    }

    if show_metadata:
        result["metadata"] = {
            "size": os.path.getsize(file_path),
            "encoding": "utf-8"
        }

    return result


# ============================================================================
# OCR — РАСПОЗНАВАНИЕ ТЕКСТА
# ============================================================================

def check_tesseract() -> bool:
    """Проверить наличие Tesseract."""
    try:
        import pytesseract
        pytesseract.pytesseract.tesseract_cmd  # Проверяем что можно импортировать
        return True
    except Exception:
        return False


def extract_ocr(file_path: str, lang: str = 'rus+eng') -> dict:
    """
    Распознать текст с изображения или PDF-скана через OCR.
    """
    try:
        from PIL import Image
        import pytesseract
    except ImportError:
        return {
            "error": "Для OCR требуется: pip install pytesseract Pillow",
            "hint": "Также установите Tesseract: https://github.com/UB-Mannheim/tesseract/wiki"
        }

    # Проверяем tesseract
    if not check_tesseract():
        return {
            "error": "Tesseract не установлен",
            "hint": "Windows: https://github.com/UB-Mannheim/tesseract/wiki\nLinux: sudo apt install tesseract-ocr"
        }

    ext = Path(file_path).suffix.lower()
    result = {
        "format": "ocr",
        "source": ext,
        "lang": lang,
        "pages": []
    }

    try:
        # Изображение (JPG, PNG, TIFF)
        if ext in ['.jpg', '.jpeg', '.png', '.tif', '.tiff', '.bmp']:
            img = Image.open(file_path)
            text = pytesseract.image_to_string(img, lang=lang)
            result["pages"].append({
                "page": 1,
                "text": text.strip()
            })

        # PDF-скан
        elif ext == '.pdf':
            import pdfplumber

            with pdfplumber.open(file_path) as pdf:
                result["page_count"] = len(pdf.pages)

                for i, page in enumerate(pdf.pages):
                    try:
                        # Конвертируем страницу PDF в изображение
                        img = page.to_image()

                        # Сохраняем во временный файл
                        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
                            tmp_path = tmp.name
                            img.save(tmp_path)

                        # OCR
                        text = pytesseract.image_to_string(Image.open(tmp_path), lang=lang)

                        result["pages"].append({
                            "page": i + 1,
                            "text": text.strip()
                        })

                        # Удаляем временный файл
                        os.unlink(tmp_path)

                    except Exception as e:
                        result["pages"].append({
                            "page": i + 1,
                            "error": str(e)
                        })

        else:
            return {
                "error": f"OCR не поддерживает формат: {ext}",
                "supported": [".jpg", ".jpeg", ".png", ".tif", ".tiff", ".bmp", ".pdf"]
            }

    except Exception as e:
        return {"error": f"Ошибка OCR: {str(e)}"}

    return result


def format_ocr_output(data: dict) -> str:
    """Форматировать результат OCR."""
    output = []
    output.append(f"{Colors.BOLD}{Colors.HEADER}=== OCR Распознавание ==={Colors.ENDC}")
    output.append(f"Источник: {data.get('source', '?')}")
    output.append(f"Язык: {data.get('lang', '?')}")

    if "page_count" in data:
        output.append(f"Страниц: {data['page_count']}")

    output.append("")

    for page in data.get("pages", []):
        if "error" in page:
            output.append(f"{Colors.RED}Страница {page['page']}: Ошибка{Colors.ENDC}")
            continue

        output.append(f"{Colors.BOLD}--- Страница {page['page']} ---{Colors.ENDC}")
        if page.get("text"):
            output.append(page["text"])
        else:
            output.append(f"{Colors.YELLOW}Текст не обнаружен{Colors.ENDC}")
        output.append("")

    return "\n".join(output)


# ============================================================================
# КОНВЕРТАЦИЯ ДОКУМЕНТОВ
# ============================================================================

def get_converter(to_format: str):
    """Получить функцию конвертера для указанного формата."""
    converters = {
        'pdf': convert_to_pdf,
        'docx': convert_to_docx,
        'md': convert_to_markdown,
        'html': convert_to_html,
    }
    return converters.get(to_format)


def check_tool(name: str) -> bool:
    """Проверить наличие утилиты в системе."""
    return shutil.which(name) is not None


def convert_to_pdf(input_path: str, output_path: str) -> dict:
    """
    Конвертировать документ в PDF.
    Поддерживает: DOCX, XLSX -> PDF через LibreOffice
    """
    # Проверяем LibreOffice
    if not check_tool('soffice') and not check_tool('libreoffice'):
        return {
            "error": "LibreOffice не установлен. Установите: https://www.libreoffice.org/",
            "hint": "Windows: скачайте установщик с официального сайта"
        }

    try:
        # LibreOffice headless mode
        cmd = [
            'soffice' if check_tool('soffice') else 'libreoffice',
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', str(Path(output_path).parent),
            input_path
        ]

        print_info("Конвертация через LibreOffice...")
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)

        if result.returncode != 0:
            return {"error": f"Ошибка конвертации: {result.stderr}"}

        # LibreOffice создаёт файл с тем же именем, но .pdf
        input_file = Path(input_path)
        expected_output = Path(output_path).parent / f"{input_file.stem}.pdf"

        # Если выходной путь не указан явно, используем результат
        if not output_path or output_path == input_path:
            output_path = str(expected_output)

        return {
            "success": True,
            "input": input_path,
            "output": output_path,
            "tool": "LibreOffice"
        }

    except subprocess.TimeoutExpired:
        return {"error": "Таймаут конвертации (120 сек)"}
    except Exception as e:
        return {"error": f"Ошибка: {str(e)}"}


def convert_to_docx(input_path: str, output_path: str) -> dict:
    """
    Конвертировать документ в DOCX.
    Поддерживает: Markdown, HTML -> DOCX через pandoc
    """
    input_ext = Path(input_path).suffix.lower()

    if input_ext not in ['.md', '.markdown', '.html', '.htm']:
        return {
            "error": f"Конвертация из {input_ext} в DOCX не поддерживается напрямую",
            "hint": "Используйте: Markdown/HTML -> PDF или сначала конвертируйте в Markdown"
        }

    if not check_tool('pandoc'):
        return {
            "error": "Pandoc не установлен. Установите: https://pandoc.org/installing.html",
            "hint": "Windows: choco install pandoc или скачайте установщик"
        }

    try:
        # Формируем выходной путь
        if not output_path or output_path == input_path:
            output_path = str(Path(input_path).with_suffix('.docx'))

        cmd = [
            'pandoc',
            input_path,
            '-o', output_path
        ]

        print_info("Конвертация через Pandoc...")
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)

        if result.returncode != 0:
            return {"error": f"Ошибка конвертации: {result.stderr}"}

        return {
            "success": True,
            "input": input_path,
            "output": output_path,
            "tool": "Pandoc"
        }

    except subprocess.TimeoutExpired:
        return {"error": "Таймаут конвертации (60 сек)"}
    except Exception as e:
        return {"error": f"Ошибка: {str(e)}"}


def convert_to_markdown(input_path: str, output_path: str) -> dict:
    """
    Конвертировать документ в Markdown.
    Поддерживает: DOCX, HTML -> MD через pandoc
    """
    input_ext = Path(input_path).suffix.lower()

    if input_ext not in ['.docx', '.doc', '.html', '.htm']:
        return {
            "error": f"Конвертация из {input_ext} в Markdown не поддерживается",
            "hint": "Используйте: DOCX -> PDF или сначала конвертируйте в DOCX"
        }

    if not check_tool('pandoc'):
        return {
            "error": "Pandoc не установлен. Установите: https://pandoc.org/installing.html",
            "hint": "Windows: choco install pandoc"
        }

    try:
        if not output_path or output_path == input_path:
            output_path = str(Path(input_path).with_suffix('.md'))

        cmd = [
            'pandoc',
            input_path,
            '-o', output_path
        ]

        print_info("Конвертация через Pandoc...")
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)

        if result.returncode != 0:
            return {"error": f"Ошибка конвертации: {result.stderr}"}

        return {
            "success": True,
            "input": input_path,
            "output": output_path,
            "tool": "Pandoc"
        }

    except subprocess.TimeoutExpired:
        return {"error": "Таймаут конвертации (60 сек)"}
    except Exception as e:
        return {"error": f"Ошибка: {str(e)}"}


def convert_to_html(input_path: str, output_path: str) -> dict:
    """
    Конвертировать документ в HTML.
    Поддерживает: DOCX, Markdown, PDF -> HTML через pandoc
    """
    if not check_tool('pandoc'):
        return {
            "error": "Pandoc не установлен. Установите: https://pandoc.org/installing.html",
            "hint": "Windows: choco install pandoc"
        }

    try:
        if not output_path or output_path == input_path:
            output_path = str(Path(input_path).with_suffix('.html'))

        cmd = [
            'pandoc',
            input_path,
            '-o', output_path
        ]

        print_info("Конвертация через Pandoc...")
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)

        if result.returncode != 0:
            return {"error": f"Ошибка конвертации: {result.stderr}"}

        return {
            "success": True,
            "input": input_path,
            "output": output_path,
            "tool": "Pandoc"
        }

    except subprocess.TimeoutExpired:
        return {"error": "Таймаут конвертации (60 сек)"}
    except Exception as e:
        return {"error": f"Ошибка: {str(e)}"}


def convert_document(input_path: str, to_format: str, output_path: str = None) -> dict:
    """
    Конвертировать документ в указанный формат.

    Args:
        input_path: Путь к исходному файлу
        to_format: Формат результата (pdf, docx, md, html)
        output_path: Путь для сохранения (опционально)

    Returns:
        dict: Результат конвертации
    """
    converter = get_converter(to_format)

    if not converter:
        return {
            "error": f"Неподдерживаемый формат: {to_format}",
            "supported": ["pdf", "docx", "md", "html"]
        }

    if not os.path.exists(input_path):
        return {"error": f"Файл не найден: {input_path}"}

    print_info(f"Конвертация {Path(input_path).suffix} -> {to_format}")
    return converter(input_path, output_path)


def format_text_output(data: dict, max_table_rows: int = 20) -> str:
    """Форматировать результат для отображения в текстовом виде."""
    output = []

    fmt = data.get("format", "unknown")

    if fmt == "pdf":
        output.append(f"{Colors.BOLD}{Colors.HEADER}=== PDF Документ ==={Colors.ENDC}")

        if "metadata" in data:
            meta = data["metadata"]
            output.append(f"Страниц: {meta.get('pages', '?')}")
            if meta.get('title'):
                output.append(f"Заголовок: {meta['title']}")
            if meta.get('author'):
                output.append(f"Автор: {meta['author']}")
            output.append("")

        for page_data in data.get("pages", []):
            output.append(f"{Colors.BOLD}--- Страница {page_data['page']} ---{Colors.ENDC}")

            if "text" in page_data:
                output.append(page_data["text"])

            if "tables" in page_data:
                for table in page_data["tables"]:
                    output.append(f"\n{Colors.YELLOW}Таблица {table['table_num']} ({table['rows']} строк):{Colors.ENDC}")
                    for i, row in enumerate(table["data"][:max_table_rows]):
                        output.append(" | ".join(str(cell) if cell else "" for cell in row))
                    if table["rows"] > max_table_rows:
                        output.append(f"... и ещё {table['rows'] - max_table_rows} строк")

    elif fmt == "docx":
        output.append(f"{Colors.BOLD}{Colors.HEADER}=== Word Документ ==={Colors.ENDC}")

        if "metadata" in data:
            meta = data["metadata"]
            if meta.get('title'):
                output.append(f"Заголовок: {meta['title']}")
            if meta.get('author'):
                output.append(f"Автор: {meta['author']}")
            output.append("")

        if data.get("paragraphs"):
            output.append(f"{Colors.BOLD}Текст ({len(data['paragraphs'])} абзацев):{Colors.ENDC}")
            for para in data["paragraphs"]:
                if para.get("style") and para["style"] != "Normal":
                    output.append(f"[{para['style']}]")
                output.append(para["text"])
                output.append("")

        if data.get("tables"):
            for table in data["tables"]:
                output.append(f"\n{Colors.YELLOW}Таблица {table['num']} ({table['rows']}x{table['cols']}):{Colors.ENDC}")
                for i, row in enumerate(table["data"][:max_table_rows]):
                    output.append(" | ".join(cell for cell in row))
                if table["rows"] > max_table_rows:
                    output.append(f"... и ещё {table['rows'] - max_table_rows} строк")

    elif fmt == "xlsx":
        output.append(f"{Colors.BOLD}{Colors.HEADER}=== Excel Файл ==={Colors.ENDC}")

        if "metadata" in data:
            meta = data["metadata"]
            if meta.get('title'):
                output.append(f"Название: {meta['title']}")
            if meta.get('author'):
                output.append(f"Автор: {meta['author']}")
            output.append("")

        output.append(f"Листов: {data.get('sheet_count', 0)}")
        output.append(f"Листы: {', '.join(data.get('sheet_names', []))}")
        output.append("")

        for sheet in data.get("sheets", []):
            output.append(f"{Colors.BOLD}--- {sheet['name']} ({sheet['max_row']} строк, {sheet['max_col']} столбцов) ---{Colors.ENDC}")

            if sheet.get("data"):
                # Заголовок таблицы (первая строка)
                header = sheet["data"][0] if sheet["data"] else []
                output.append(" | ".join(str(cell) if cell else "" for cell in header))
                output.append("-" * 50)

                # Данные
                for i, row in enumerate(sheet["data"][:max_table_rows]):
                    output.append(" | ".join(str(cell) if cell else "" for cell in row))

                if sheet.get("truncated"):
                    output.append(f"\n... показано {sheet['preview_rows']} из {sheet['total_rows']} строк")

    elif fmt == "txt":
        output.append(f"{Colors.BOLD}{Colors.HEADER}=== Текстовый файл ==={Colors.ENDC}")
        output.append(f"Строк: {data.get('lines', 0)}, Символов: {data.get('chars', 0)}")
        output.append("")
        output.append(data.get("text", ""))

    return "\n".join(output)


def main():
    parser = argparse.ArgumentParser(
        description="Извлечение текста, конвертация и OCR",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Примеры использования:
  # Извлечение текста
  python extract_text.py document.pdf
  python extract_text.py document.docx --metadata
  python extract_text.py spreadsheet.xlsx --tables

  # Конвертация
  python extract_text.py document.docx --convert --to pdf
  python extract_text.py document.docx --convert --to md
  python extract_text.py file.md --convert --to docx

  # OCR
  python extract_text.py scan.jpg --ocr --lang rus+eng
  python extract_text.py scanned.pdf --ocr
        """
    )

    parser.add_argument("file", help="Путь к файлу")
    parser.add_argument("--tables", "-t", action="store_true",
                        help="Извлекать таблицы")
    parser.add_argument("--metadata", "-m", action="store_true",
                        help="Показать метаданные")
    parser.add_argument("--output", "-o",
                        help="Сохранить результат в файл")
    parser.add_argument("--json", "-j", action="store_true",
                        help="Вывод в формате JSON")

    # Конвертация
    parser.add_argument("--convert", "-c", action="store_true",
                        help="Конвертировать файл")
    parser.add_argument("--to", choices=["pdf", "docx", "md", "html"],
                        help="Формат для конвертации")

    # OCR
    parser.add_argument("--ocr", action="store_true",
                        help="Распознать текст (OCR)")
    parser.add_argument("--lang", default="rus+eng",
                        help="Язык OCR (rus, eng, rus+eng)")

    args = parser.parse_args()

    # Проверка файла
    if not os.path.exists(args.file):
        print_error(f"Файл не найден: {args.file}")
        sys.exit(1)

    file_size = os.path.getsize(args.file)
    print_info(f"Файл: {args.file} ({file_size / 1024:.1f} KB)")

    # Режим OCR
    if args.ocr:
        print_info("Режим: OCR распознавание")
        data = extract_ocr(args.file, args.lang)

        if "error" in data:
            print_error(data["error"])
            if "hint" in data:
                print_warning(data["hint"])
            sys.exit(1)

        if args.json:
            print(json.dumps(data, ensure_ascii=False, indent=2, default=str))
        else:
            print(format_ocr_output(data))

        if args.output:
            with open(args.output, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2, default=str)
            print_success(f"Сохранено: {args.output}")
        return

    # Режим конвертации
    if args.convert:
        if not args.to:
            print_error("Укажите формат конвертации: --to pdf|docx|md|html")
            sys.exit(1)

        print_info(f"Режим: Конвертация в {args.to.upper()}")
        result = convert_document(args.file, args.to, args.output)

        if "error" in result:
            print_error(result["error"])
            if "hint" in result:
                print_warning(result["hint"])
            sys.exit(1)

        if result.get("success"):
            print_success(f"Конвертация завершена!")
            print_info(f"Инструмент: {result.get('tool')}")
            print_info(f"Результат: {result.get('output')}")
        return

    # Режим извлечения текста
    fmt = detect_format(args.file)

    if fmt == "unknown":
        print_error(f"Неподдерживаемый формат файла: {args.file}")
        sys.exit(1)

    print_info(f"Формат: {fmt.upper()}")
    print_info("Извлечение...")

    # Извлечение данных
    try:
        if fmt == "pdf":
            data = extract_pdf(args.file, args.tables, args.metadata)
        elif fmt == "docx":
            data = extract_docx(args.file, args.metadata)
        elif fmt == "xlsx":
            data = extract_xlsx(args.file, args.tables, args.metadata)
        elif fmt == "txt":
            data = extract_txt(args.file, args.metadata)
        else:
            data = {"error": f"Формат {fmt} не поддерживается"}
    except Exception as e:
        print_error(f"Ошибка при извлечении: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

    # Обработка ошибок
    if "error" in data:
        print_error(data["error"])
        sys.exit(1)

    # Вывод
    if args.json:
        print(json.dumps(data, ensure_ascii=False, indent=2, default=str))
    else:
        print(format_text_output(data))

    # Сохранение в файл
    if args.output:
        with open(args.output, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2, default=str)
        print_success(f"Сохранено: {args.output}")


if __name__ == "__main__":
    main()
